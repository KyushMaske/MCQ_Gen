from groq import Groq
import os
from docx import Document
from dotenv import load_dotenv

import fitz  
from docx import Document
from io import BytesIO

load_dotenv()

api_key = os.getenv("GROQ_API_KEY")

client = Groq(api_key=api_key)


def read_pdf(file):
    """
    Extract text from a PDF file (in memory).
    """
    pdf_content = BytesIO(file.file.read())
    try:
        with pdfplumber.open(pdf_content) as pdf:
            text = "\n".join(
                page.extract_text() for page in pdf.pages if page.extract_text()
            )
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Error reading PDF: {str(e)}")
    return text


def read_docx(file):
    """
    Extract text from a DOCX file (in memory).
    """
    docx_content = BytesIO(file.file.read())
    try:
        document = Document(docx_content)
        text = "\n".join(para.text for para in document.paragraphs)
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Error reading DOCX: {str(e)}")
    return text


def read_txt(file):
    """
    Extract text from a TXT file (in memory).
    """
    try:
        text = file.file.read().decode("utf-8")
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Error reading TXT: {str(e)}")
    return text


def get_response(messages):
    try:
        stream = client.chat.completions.create(
            messages=messages,
            model="llama3-8b-8192",  # Adjust to your model
            temperature=0.5,
            max_tokens=1024,
            top_p=1,
            stop=None,
            stream=True,
        )

        response_content = ""
        for chunk in stream:
            if chunk.choices and chunk.choices[0].delta.content:
                response_content += chunk.choices[0].delta.content

        print(response_content, "response_content")
        return response_content
    except Exception as e:
        return f"Error: {str(e)}"


def generate_mcqs(
    text,
    num_questions,
    difficulty_level,
    include_topics,
    exclude_topics,
    question_format,
):
    prompt = f"Generate {num_questions} {difficulty_level} multiple choice questions (MCQs) with options (A, B, C, D) based on the following text:\n{text}"

    if include_topics:
        prompt += f"\nInclude these topics: {include_topics}"
    if exclude_topics:
        prompt += f"\nExclude these topics: {exclude_topics}"
    if question_format == "true_false":
        prompt = prompt.replace("multiple choice questions", "true/false questions")

    return get_response(
        [
            {
                "role": "system",
                "content": "You are an assistant that generates multiple choice or true/false questions from a given text.",
            },
            {"role": "user", "content": prompt},
        ]
    )


def save_mcqs_to_doc(mcqs, file_name="mcqs.docx"):
    """Saves the generated MCQs to a .docx file."""
    SAVE_DIR = "generated_docs"
    os.makedirs(SAVE_DIR, exist_ok=True)

    file_path = os.path.join(SAVE_DIR, file_name)
    document = Document()

    document.add_heading("Generated MCQs", level=1)
    for i, mcq in enumerate(mcqs.split("\n\n"), start=1):
        document.add_paragraph(f"Q{i}. {mcq.strip()}", style="List Number")

    document.save(file_path)
    return file_path
